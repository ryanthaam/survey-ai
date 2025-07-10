import pandas as pd
import numpy as np
import io
import tempfile
import os
from typing import List, Dict, Tuple, Optional, Union
import streamlit as st

# PDF processing
try:
    import PyPDF2
    import pdfplumber
except ImportError:
    PyPDF2 = None
    pdfplumber = None

# Excel processing
try:
    import openpyxl
    import xlrd
except ImportError:
    openpyxl = None
    xlrd = None

# Word processing
try:
    import docx
except ImportError:
    docx = None

class FileProcessor:
    """Universal file processor for various survey data formats."""
    
    def __init__(self):
        self.supported_formats = {
            'csv': ['csv'],
            'excel': ['xlsx', 'xls'],
            'pdf': ['pdf'],
            'word': ['docx', 'doc'],
            'text': ['txt']
        }
        
    def get_file_type(self, uploaded_file) -> str:
        """Determine file type from uploaded file."""
        if uploaded_file is None:
            return 'unknown'
        
        # Get file extension
        file_extension = uploaded_file.name.split('.')[-1].lower()
        
        # Map to category
        for category, extensions in self.supported_formats.items():
            if file_extension in extensions:
                return category
        
        return 'unknown'
    
    def process_file(self, uploaded_file) -> Tuple[Optional[pd.DataFrame], List[str], Dict]:
        """Process uploaded file and return DataFrame and/or text data."""
        file_type = self.get_file_type(uploaded_file)
        
        try:
            if file_type == 'csv':
                return self._process_csv(uploaded_file)
            elif file_type == 'excel':
                return self._process_excel(uploaded_file)
            elif file_type == 'pdf':
                return self._process_pdf(uploaded_file)
            elif file_type == 'word':
                return self._process_word(uploaded_file)
            elif file_type == 'text':
                return self._process_text(uploaded_file)
            else:
                return None, [], {'error': f'Unsupported file type: {file_type}'}
                
        except Exception as e:
            return None, [], {'error': f'Error processing {file_type} file: {str(e)}'}
    
    def _process_csv(self, uploaded_file) -> Tuple[pd.DataFrame, List[str], Dict]:
        """Process CSV files."""
        try:
            df = pd.read_csv(uploaded_file)
            info = {
                'file_type': 'csv',
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': df.columns.tolist()
            }
            return df, [], info
        except Exception as e:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252', 'iso-8859-1']
            for encoding in encodings:
                try:
                    uploaded_file.seek(0)
                    df = pd.read_csv(uploaded_file, encoding=encoding)
                    info = {
                        'file_type': 'csv',
                        'rows': len(df),
                        'columns': len(df.columns),
                        'column_names': df.columns.tolist(),
                        'encoding_used': encoding
                    }
                    return df, [], info
                except:
                    continue
            
            raise Exception(f"Could not read CSV file with any encoding: {str(e)}")
    
    def _process_excel(self, uploaded_file) -> Tuple[pd.DataFrame, List[str], Dict]:
        """Process Excel files (.xlsx, .xls)."""
        if openpyxl is None or xlrd is None:
            raise Exception("Excel processing libraries not installed. Install openpyxl and xlrd.")
        
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix='.xlsx') as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        try:
            # Try reading with pandas
            excel_file = pd.ExcelFile(tmp_path)
            
            # Get sheet names
            sheet_names = excel_file.sheet_names
            
            # If multiple sheets, ask user or take first one
            if len(sheet_names) > 1:
                st.info(f"📊 Found {len(sheet_names)} sheets: {', '.join(sheet_names)}. Using first sheet: '{sheet_names[0]}'")
            
            # Read the first sheet (or specified sheet)
            df = pd.read_excel(tmp_path, sheet_name=sheet_names[0])
            
            info = {
                'file_type': 'excel',
                'rows': len(df),
                'columns': len(df.columns),
                'column_names': df.columns.tolist(),
                'sheet_names': sheet_names,
                'sheet_used': sheet_names[0]
            }
            
            return df, [], info
            
        finally:
            # Clean up temp file
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    def _process_pdf(self, uploaded_file) -> Tuple[None, List[str], Dict]:
        """Process PDF files and extract text."""
        if PyPDF2 is None or pdfplumber is None:
            raise Exception("PDF processing libraries not installed. Install PyPDF2 and pdfplumber.")
        
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        try:
            text_blocks = []
            
            # Try pdfplumber first (better for structured data)
            try:
                with pdfplumber.open(tmp_path) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        # Try to extract tables first
                        tables = page.extract_tables()
                        if tables:
                            for table in tables:
                                # Convert table to text
                                for row in table:
                                    if row:
                                        row_text = ' | '.join([str(cell) if cell else '' for cell in row])
                                        if row_text.strip():
                                            text_blocks.append(row_text.strip())
                        
                        # Extract regular text
                        text = page.extract_text()
                        if text:
                            # Split into meaningful chunks
                            lines = [line.strip() for line in text.split('\n') if line.strip()]
                            text_blocks.extend(lines)
            
            except Exception as e:
                st.warning(f"pdfplumber failed: {str(e)}. Trying PyPDF2...")
                
                # Fallback to PyPDF2
                with open(tmp_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        text = page.extract_text()
                        
                        if text:
                            lines = [line.strip() for line in text.split('\n') if line.strip()]
                            text_blocks.extend(lines)
            
            # Filter meaningful text
            meaningful_texts = []
            for text in text_blocks:
                if len(text) > 10 and not text.isdigit():  # Filter out page numbers, etc.
                    meaningful_texts.append(text)
            
            info = {
                'file_type': 'pdf',
                'total_text_blocks': len(text_blocks),
                'meaningful_texts': len(meaningful_texts),
                'extraction_method': 'pdfplumber + PyPDF2'
            }
            
            return None, meaningful_texts, info
            
        finally:
            # Clean up temp file
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    def _process_word(self, uploaded_file) -> Tuple[None, List[str], Dict]:
        """Process Word documents (.docx)."""
        if docx is None:
            raise Exception("Word processing library not installed. Install python-docx.")
        
        # Save uploaded file temporarily
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp_file:
            tmp_file.write(uploaded_file.getvalue())
            tmp_path = tmp_file.name
        
        try:
            doc = docx.Document(tmp_path)
            
            text_blocks = []
            
            # Extract paragraphs
            for paragraph in doc.paragraphs:
                text = paragraph.text.strip()
                if text and len(text) > 5:
                    text_blocks.append(text)
            
            # Extract from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    
                    if row_text:
                        combined_text = ' | '.join(row_text)
                        if len(combined_text) > 10:
                            text_blocks.append(combined_text)
            
            info = {
                'file_type': 'word',
                'paragraphs': len(doc.paragraphs),
                'tables': len(doc.tables),
                'text_blocks_extracted': len(text_blocks)
            }
            
            return None, text_blocks, info
            
        finally:
            # Clean up temp file
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    def _process_text(self, uploaded_file) -> Tuple[None, List[str], Dict]:
        """Process plain text files."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    uploaded_file.seek(0)
                    content = uploaded_file.read().decode(encoding)
                    break
                except UnicodeDecodeError:
                    continue
            else:
                raise Exception("Could not decode text file with any standard encoding")
            
            # Split into meaningful lines
            lines = [line.strip() for line in content.split('\n') if line.strip()]
            
            # Filter out very short lines
            meaningful_lines = [line for line in lines if len(line) > 5]
            
            info = {
                'file_type': 'text',
                'total_lines': len(lines),
                'meaningful_lines': len(meaningful_lines),
                'encoding_used': encoding
            }
            
            return None, meaningful_lines, info
            
        except Exception as e:
            raise Exception(f"Error processing text file: {str(e)}")
    
    def convert_text_to_dataframe(self, texts: List[str], method: str = 'auto') -> pd.DataFrame:
        """Convert extracted text to a DataFrame for analysis."""
        
        if method == 'auto':
            # Try to detect structure
            if self._looks_like_structured_data(texts):
                return self._parse_structured_text(texts)
            else:
                return self._create_simple_dataframe(texts)
        elif method == 'simple':
            return self._create_simple_dataframe(texts)
        elif method == 'structured':
            return self._parse_structured_text(texts)
        else:
            return self._create_simple_dataframe(texts)
    
    def _looks_like_structured_data(self, texts: List[str]) -> bool:
        """Check if text looks like structured data (tables, Q&A, etc.)."""
        structured_indicators = 0
        
        for text in texts[:10]:  # Check first 10 lines
            # Look for common structured patterns
            if any(pattern in text.lower() for pattern in ['question', 'q:', 'answer', 'a:', '|', '\t']):
                structured_indicators += 1
            if text.count('|') > 1:  # Table-like structure
                structured_indicators += 2
            if text.startswith(('1.', '2.', '3.', 'Q1', 'Q2', 'Q3')):  # Numbered items
                structured_indicators += 1
        
        return structured_indicators > 3
    
    def _parse_structured_text(self, texts: List[str]) -> pd.DataFrame:
        """Parse structured text into columns."""
        rows = []
        
        for text in texts:
            # Try to split by common delimiters
            if '|' in text:
                # Table-like structure
                parts = [part.strip() for part in text.split('|')]
                rows.append(parts)
            elif ':' in text and len(text.split(':')) == 2:
                # Question: Answer format
                question, answer = text.split(':', 1)
                rows.append([question.strip(), answer.strip()])
            else:
                # Single column
                rows.append([text])
        
        # Create DataFrame with appropriate columns
        if rows:
            max_cols = max(len(row) for row in rows)
            
            # Pad rows to same length
            for row in rows:
                while len(row) < max_cols:
                    row.append('')
            
            # Create column names
            if max_cols == 1:
                columns = ['Text']
            elif max_cols == 2:
                columns = ['Question', 'Response']
            else:
                columns = [f'Column_{i+1}' for i in range(max_cols)]
            
            df = pd.DataFrame(rows, columns=columns)
            return df
        else:
            return pd.DataFrame({'Text': texts})
    
    def _create_simple_dataframe(self, texts: List[str]) -> pd.DataFrame:
        """Create a simple DataFrame with text in one column."""
        return pd.DataFrame({'Text': texts})
    
    def get_supported_formats(self) -> Dict[str, List[str]]:
        """Get dictionary of supported file formats."""
        return self.supported_formats.copy()
    
    def is_format_supported(self, file_extension: str) -> bool:
        """Check if a file format is supported."""
        file_extension = file_extension.lower().lstrip('.')
        
        for extensions in self.supported_formats.values():
            if file_extension in extensions:
                return True
        
        return False